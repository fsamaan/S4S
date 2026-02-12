"""Export utilities for clinical research analysis."""

import json
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import CellIsRule

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


def sanitize_sheet_name(name: str) -> str:
    """Sanitize Excel sheet names to avoid invalid characters."""
    invalid = set(r"[]:*?/\\")
    cleaned = "".join(ch for ch in name if ch not in invalid).strip()
    if not cleaned:
        cleaned = "Sheet"
    return cleaned[:31]


def excel_safe_value(value: Any) -> Any:
    """Ensure values are Excel-compatible."""
    if isinstance(value, (list, tuple, dict, set)):
        return str(value)
    return value


def to_excel_bytes(dict_of_sheets: Dict[str, Any]) -> bytes:
    """Create an Excel workbook from named sheets and return bytes."""
    wb = Workbook()
    first = True

    for sheet_name, content in dict_of_sheets.items():
        safe_name = sanitize_sheet_name(sheet_name)
        if first:
            ws = wb.active
            ws.title = safe_name
            first = False
        else:
            ws = wb.create_sheet(title=safe_name)

        if isinstance(content, pd.DataFrame):
            start_row = 1
            title_text = getattr(content, "attrs", {}).get("title")
            if title_text:
                ws.cell(row=1, column=1, value=title_text).font = Font(bold=True)
                start_row = 3
            for r_idx, row in enumerate([content.columns.tolist()] + content.values.tolist(), start=start_row):
                for c_idx, value in enumerate(row, start=1):
                    ws.cell(row=r_idx, column=c_idx, value=excel_safe_value(value))

            header_row = start_row
            for cell in ws[header_row]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center")

            ws.freeze_panes = f"A{header_row + 1}"

            for col_idx, col in enumerate(content.columns, start=1):
                max_len = max(
                    [len(str(col))] + [len(str(x)) for x in content.iloc[:, col_idx - 1].astype(str).tolist()]
                )
                ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 2, 40)

            if "SMD" in content.columns:
                smd_col = content.columns.get_loc("SMD") + 1
                ws.conditional_formatting.add(
                    f"{get_column_letter(smd_col)}2:{get_column_letter(smd_col)}{len(content) + 1}",
                    CellIsRule(operator=">", formula=["0.1"], font=Font(bold=True)),
                )

            if "p-value" in content.columns:
                p_col = content.columns.get_loc("p-value") + 1
                ws.conditional_formatting.add(
                    f"{get_column_letter(p_col)}2:{get_column_letter(p_col)}{len(content) + 1}",
                    CellIsRule(operator="<", formula=["0.05"], font=Font(bold=True)),
                )
        else:
            ws["A1"] = str(content)

    stream = BytesIO()
    wb.save(stream)
    return stream.getvalue()


def round_numeric_columns(df: pd.DataFrame, decimals: int = 3) -> pd.DataFrame:
    """Round numeric columns in a DataFrame."""
    rounded = df.copy()
    numeric_cols = rounded.select_dtypes(include="number").columns
    if len(numeric_cols) > 0:
        rounded[numeric_cols] = rounded[numeric_cols].round(decimals)
    return rounded


def clean_regression_results(df: pd.DataFrame) -> pd.DataFrame:
    """Clean regression results DataFrame for export."""
    cleaned = df.copy()
    if "index" in cleaned.columns:
        cleaned = cleaned.drop(columns=["index"])
    if cleaned.index.name is not None or not isinstance(cleaned.index, pd.RangeIndex):
        cleaned = cleaned.reset_index(drop=True)
    return cleaned


def build_run_config_json(config: Dict[str, Any]) -> str:
    """Build run configuration JSON string."""
    return json.dumps(config, indent=2, default=str)


def create_methods_text(
    table1_opts: Dict[str, Any],
    table2_opts: Dict[str, Any],
    model_opts: Dict[str, Any],
    missing_strategy: str,
) -> str:
    """Generate a concise editable methods text block."""
    base = (
        f"Baseline characteristics were summarized as {table1_opts.get('summary', 'mean (SD) or median (IQR)')} "
        f"and compared using {table1_opts.get('tests', 'appropriate parametric or nonparametric tests')}; "
        f"standardized mean differences were calculated to assess balance. "
        f"Crude outcome comparisons used {table2_opts.get('crude', 'unadjusted regression models')} by exposure group. "
        f"Adjusted models included covariates selected a priori and from imbalance diagnostics. "
        f"Missing data handling: {missing_strategy}."
    )
    return base


def add_docx_table(doc: "Document", title: str, df: pd.DataFrame) -> None:
    """Add a table to a Word document."""
    if not DOCX_AVAILABLE:
        return
    doc.add_heading(title, level=2)
    if df is None or df.empty:
        doc.add_paragraph("No data available.")
        return
    if df.index.name is not None or not isinstance(df.index, pd.RangeIndex):
        df = df.reset_index()
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else str(val)


def generate_word_report(
    methods_text: str,
    table1: Optional[pd.DataFrame],
    table2: Optional[pd.DataFrame],
    table2_or: Optional[pd.DataFrame],
    model_results: Optional[pd.DataFrame],
    model_compare: Optional[pd.DataFrame],
    stratified_table2: Optional[Dict[str, pd.DataFrame]],
    custom_summary_table: Optional[pd.DataFrame],
) -> Optional[bytes]:
    """Generate a Word report with all tables."""
    if not DOCX_AVAILABLE:
        return None

    doc = Document()
    doc.add_heading("Analysis Report", level=1)

    if methods_text and methods_text.strip():
        doc.add_heading("Methods", level=2)
        doc.add_paragraph(methods_text)

    if table1 is not None:
        add_docx_table(doc, "Table 1", table1)

    if table2 is not None:
        add_docx_table(doc, "Table 2", table2)

    if table2_or is not None:
        add_docx_table(doc, "Table 2 (OR)", table2_or)

    if model_results is not None:
        add_docx_table(
            doc,
            "Regression Results",
            round_numeric_columns(clean_regression_results(model_results)),
        )

    if model_compare is not None:
        add_docx_table(doc, "Model Comparison", round_numeric_columns(model_compare))

    if stratified_table2:
        for level, table in stratified_table2.items():
            add_docx_table(doc, f"Stratified Table 2: {level}", table)

    if custom_summary_table is not None:
        add_docx_table(doc, "Custom Summary Table", custom_summary_table)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
